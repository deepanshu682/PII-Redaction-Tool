"""
docx_redactor.py - Structure-Preserving DOCX Redactor
Redacts PII across paragraphs, tables, headers, and footers in Microsoft Word documents.
"""

import docx
from pii_redactor import PIIRedactor, PIIItem
from typing import Dict, List, Tuple

class DOCXRedactor:
    def __init__(self, redactor: PIIRedactor):
        self.redactor = redactor
        self.stats: Dict[str, int] = {}
        self.total_replacements = 0

    def _redact_paragraph(self, p: docx.text.paragraph.Paragraph) -> int:
        """Redact PII in a single paragraph while preserving styling."""
        if not p.text or not p.text.strip():
            return 0

        original_text = p.text
        redacted_text, entities = self.redactor.redact_text(original_text)

        if not entities or redacted_text == original_text:
            return 0

        # Update stats
        for ent in entities:
            self.stats[ent.entity_type] = self.stats.get(ent.entity_type, 0) + 1
            self.total_replacements += 1

        # Preserve formatting by setting text on the first run and clearing remaining runs
        if p.runs:
            p.runs[0].text = redacted_text
            for run in p.runs[1:]:
                run.text = ""
        else:
            p.text = redacted_text

        return len(entities)

    def redact_document(self, input_docx_path: str, output_docx_path: str) -> Dict[str, int]:
        """Read DOCX, redact PII in all sections, and save to output_docx_path."""
        doc = docx.Document(input_docx_path)
        self.stats = {}
        self.total_replacements = 0

        # 1. Redact Body Paragraphs
        for p in doc.paragraphs:
            self._redact_paragraph(p)

        # 2. Redact Tables (cells & nested paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        self._redact_paragraph(p)

        # 3. Redact Headers and Footers
        for section in doc.sections:
            if section.header:
                for p in section.header.paragraphs:
                    self._redact_paragraph(p)
                for table in section.header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                self._redact_paragraph(p)

            if section.footer:
                for p in section.footer.paragraphs:
                    self._redact_paragraph(p)
                for table in section.footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                self._redact_paragraph(p)

        # Save redacted document
        doc.save(output_docx_path)
        return self.stats

if __name__ == "__main__":
    redactor = PIIRedactor(mode="SYNTHETIC")
    docx_redactor = DOCXRedactor(redactor)
    stats = docx_redactor.redact_document("Red Herring Prospectus.docx", "Red_Herring_Prospectus_Redacted.docx")
    print("\n=== DOCX REDACTION COMPLETE ===")
    print(f"Total PII Replacements: {docx_redactor.total_replacements}")
    print("Breakdown by entity type:")
    for etype, count in stats.items():
        print(f" - {etype}: {count}")
